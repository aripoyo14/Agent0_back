# app/services/file_analyzer_full.py
"""
ファイル内容解析サービス（完全版）
- PDF、Word、Excel、テキストファイルの内容を完全抽出
- AI返信生成時にファイル内容を詳細に考慮
- 処理時間は長くなるが、より詳細な分析が可能
"""

import requests
import io
import time
from typing import Optional, Dict, Any, List
from fastapi import HTTPException
import PyPDF2
from docx import Document
from openpyxl import load_workbook
import re
import json

def extract_file_content_full(file_url: str, file_type: str, file_name: str) -> Dict[str, Any]:
    """
    ファイルURLから内容を完全抽出する
    完全版：ファイル全体を読み取り、構造化されたテキストとして抽出
    """
    start_time = time.time()
    
    try:
        # ファイルをダウンロード
        response = requests.get(file_url, timeout=30)  # タイムアウトを30秒に延長
        response.raise_for_status()
        file_content = response.content
        
        # ファイルサイズ制限（10MBまで）
        MAX_SIZE = 10 * 1024 * 1024
        if len(file_content) > MAX_SIZE:
            return {
                "file_name": file_name,
                "file_type": file_type,
                "content": "ファイルサイズが大きすぎます（10MB制限）",
                "summary": "ファイルサイズ制限超過",
                "processing_time": time.time() - start_time,
                "file_size_mb": len(file_content) / (1024 * 1024)
            }
        
        # ファイルタイプ別の完全解析
        if file_type == "application/pdf":
            result = extract_pdf_content_full(file_content, file_name)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            result = extract_word_content_full(file_content, file_name)
        elif file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
            result = extract_excel_content_full(file_content, file_name)
        elif file_type.startswith("text/"):
            result = extract_text_content_full(file_content, file_name)
        else:
            result = {
                "file_name": file_name,
                "file_type": file_type,
                "content": "このファイル形式は現在サポートされていません。",
                "summary": "未対応ファイル形式",
                "structure": {},
                "key_points": []
            }
        
        result["processing_time"] = time.time() - start_time
        result["file_size_mb"] = len(file_content) / (1024 * 1024)
        return result
            
    except Exception as e:
        return {
            "file_name": file_name,
            "file_type": file_type,
            "content": f"ファイル読み取りエラー: {str(e)}",
            "summary": "読み取り失敗",
            "processing_time": time.time() - start_time,
            "error": str(e)
        }

def extract_pdf_content_full(file_content: bytes, file_name: str) -> Dict[str, Any]:
    """PDFファイルの内容を完全抽出"""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # 全ページのテキストを抽出
        full_text = ""
        page_texts = []
        structure = {
            "total_pages": len(pdf_reader.pages),
            "pages": []
        }
        
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            page_texts.append(page_text)
            full_text += f"\n--- ページ {i+1} ---\n{page_text}"
            
            # ページ構造を分析
            page_structure = {
                "page_number": i + 1,
                "text_length": len(page_text),
                "has_tables": "表" in page_text or "TABLE" in page_text.upper(),
                "has_figures": "図" in page_text or "FIGURE" in page_text.upper(),
                "key_sections": extract_key_sections(page_text)
            }
            structure["pages"].append(page_structure)
        
        # キーポイントを抽出
        key_points = extract_key_points_from_text(full_text)
        
        # 概要を生成（最初の500文字）
        summary = full_text[:500] + "..." if len(full_text) > 500 else full_text
        
        return {
            "file_name": file_name,
            "file_type": "application/pdf",
            "content": full_text,
            "summary": summary,
            "structure": structure,
            "key_points": key_points,
            "total_pages": len(pdf_reader.pages),
            "total_text_length": len(full_text)
        }
    except Exception as e:
        return {
            "file_name": file_name,
            "file_type": "application/pdf",
            "content": f"PDF読み取りエラー: {str(e)}",
            "summary": "PDF読み取り失敗",
            "structure": {},
            "key_points": []
        }

def extract_word_content_full(file_content: bytes, file_name: str) -> Dict[str, Any]:
    """Wordファイルの内容を完全抽出"""
    try:
        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)
        
        # 全段落とテーブルからテキストを抽出
        full_text = ""
        paragraphs = []
        tables = []
        structure = {
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "sections": []
        }
        
        # 段落を抽出
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append(para.text)
                full_text += para.text + "\n"
        
        # テーブルを抽出
        for i, table in enumerate(doc.tables):
            table_text = f"\n--- テーブル {i+1} ---\n"
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_text += " | ".join(row_data) + "\n"
            tables.append(table_text)
            full_text += table_text
        
        # 構造を分析
        structure["sections"] = extract_document_sections(full_text)
        
        # キーポイントを抽出
        key_points = extract_key_points_from_text(full_text)
        
        # 概要を生成
        summary = full_text[:500] + "..." if len(full_text) > 500 else full_text
        
        return {
            "file_name": file_name,
            "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "content": full_text,
            "summary": summary,
            "structure": structure,
            "key_points": key_points,
            "paragraphs": paragraphs,
            "tables": tables,
            "total_text_length": len(full_text)
        }
    except Exception as e:
        return {
            "file_name": file_name,
            "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "content": f"Word読み取りエラー: {str(e)}",
            "summary": "Word読み取り失敗",
            "structure": {},
            "key_points": []
        }

def extract_excel_content_full(file_content: bytes, file_name: str) -> Dict[str, Any]:
    """Excelファイルの内容を完全抽出"""
    try:
        excel_file = io.BytesIO(file_content)
        wb = load_workbook(excel_file, data_only=True)
        
        full_text = ""
        sheets_data = []
        structure = {
            "total_sheets": len(wb.sheetnames),
            "sheets": []
        }
        
        # 全シートを処理
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_text = f"\n=== シート: {sheet_name} ===\n"
            sheet_data = {
                "sheet_name": sheet_name,
                "max_row": ws.max_row,
                "max_column": ws.max_column,
                "data": []
            }
            
            # シート内の全データを抽出
            for row in ws.iter_rows():
                row_data = []
                for cell in row:
                    if cell.value is not None:
                        row_data.append(str(cell.value))
                    else:
                        row_data.append("")
                if any(cell != "" for cell in row_data):
                    sheet_data["data"].append(row_data)
                    sheet_text += " | ".join(row_data) + "\n"
            
            sheets_data.append(sheet_data)
            full_text += sheet_text
            structure["sheets"].append({
                "name": sheet_name,
                "rows": ws.max_row,
                "columns": ws.max_column,
                "has_data": len(sheet_data["data"]) > 0
            })
        
        # キーポイントを抽出
        key_points = extract_key_points_from_text(full_text)
        
        # 概要を生成
        summary = f"Excelファイル: {len(wb.sheetnames)}シート、合計{sum(len(sheet['data']) for sheet in sheets_data)}行のデータ"
        
        return {
            "file_name": file_name,
            "file_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "content": full_text,
            "summary": summary,
            "structure": structure,
            "key_points": key_points,
            "sheets_data": sheets_data,
            "total_sheets": len(wb.sheetnames),
            "total_rows": sum(len(sheet['data']) for sheet in sheets_data)
        }
    except Exception as e:
        return {
            "file_name": file_name,
            "file_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "content": f"Excel読み取りエラー: {str(e)}",
            "summary": "Excel読み取り失敗",
            "structure": {},
            "key_points": []
        }

def extract_text_content_full(file_content: bytes, file_name: str) -> Dict[str, Any]:
    """テキストファイルの内容を完全抽出"""
    try:
        text_content = file_content.decode('utf-8', errors='ignore')
        
        # 構造を分析
        structure = {
            "total_lines": len(text_content.split('\n')),
            "total_words": len(text_content.split()),
            "total_characters": len(text_content),
            "sections": extract_document_sections(text_content)
        }
        
        # キーポイントを抽出
        key_points = extract_key_points_from_text(text_content)
        
        # 概要を生成
        summary = text_content[:500] + "..." if len(text_content) > 500 else text_content
        
        return {
            "file_name": file_name,
            "file_type": "text/plain",
            "content": text_content,
            "summary": summary,
            "structure": structure,
            "key_points": key_points,
            "total_text_length": len(text_content)
        }
    except Exception as e:
        return {
            "file_name": file_name,
            "file_type": "text/plain",
            "content": f"テキスト読み取りエラー: {str(e)}",
            "summary": "テキスト読み取り失敗",
            "structure": {},
            "key_points": []
        }

def extract_key_sections(text: str) -> List[str]:
    """テキストから主要セクションを抽出"""
    sections = []
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        # 見出しっぽい行を検出
        if (len(line) < 100 and 
            (line.isupper() or 
             line.startswith(('第', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) or
             '章' in line or '節' in line or '項' in line)):
            sections.append(line)
    
    return sections[:10]  # 最大10個まで

def extract_document_sections(text: str) -> List[Dict[str, Any]]:
    """ドキュメントの構造を分析"""
    sections = []
    lines = text.split('\n')
    current_section = None
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # セクション開始を検出
        if (len(line) < 100 and 
            (line.isupper() or 
             line.startswith(('第', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) or
             '章' in line or '節' in line or '項' in line)):
            
            if current_section:
                sections.append(current_section)
            
            current_section = {
                "title": line,
                "start_line": i,
                "content_length": 0
            }
        elif current_section:
            current_section["content_length"] += len(line)
    
    if current_section:
        sections.append(current_section)
    
    return sections

def extract_key_points_from_text(text: str) -> List[str]:
    """テキストからキーポイントを抽出"""
    key_points = []
    
    # 重要なキーワードを含む文を抽出
    important_keywords = [
        '重要', '要点', '結論', '結果', '提案', '推奨', '課題', '問題',
        '予算', 'コスト', '費用', 'スケジュール', '期限', '目標',
        '技術', '実装', '導入', '運用', '保守'
    ]
    
    sentences = re.split(r'[。！？\n]', text)
    
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) > 10 and len(sentence) < 200:  # 適切な長さの文
            for keyword in important_keywords:
                if keyword in sentence:
                    key_points.append(sentence)
                    break
    
    return key_points[:20]  # 最大20個まで

def compare_analysis_results(lightweight_result: Dict[str, Any], full_result: Dict[str, Any]) -> Dict[str, Any]:
    """軽量版と完全版の結果を比較"""
    return {
        "lightweight": {
            "processing_time": lightweight_result.get("processing_time", 0),
            "content_length": len(lightweight_result.get("content", "")),
            "summary_length": len(lightweight_result.get("summary", "")),
            "has_structure": "structure" not in lightweight_result,
            "has_key_points": "key_points" not in lightweight_result
        },
        "full": {
            "processing_time": full_result.get("processing_time", 0),
            "content_length": len(full_result.get("content", "")),
            "summary_length": len(full_result.get("summary", "")),
            "has_structure": "structure" in full_result,
            "has_key_points": "key_points" in full_result,
            "structure_details": full_result.get("structure", {}),
            "key_points_count": len(full_result.get("key_points", []))
        },
        "comparison": {
            "time_ratio": full_result.get("processing_time", 0) / max(lightweight_result.get("processing_time", 1), 1),
            "content_ratio": len(full_result.get("content", "")) / max(len(lightweight_result.get("content", "")), 1),
            "improvement_factor": len(full_result.get("key_points", [])) / max(len(lightweight_result.get("key_points", [])), 1)
        }
    } 